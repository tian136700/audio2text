#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将识别文本转换为 PDF 或 Word 格式
支持格式调整：说话人和时间放在上面一行，说话内容放在下面一行
"""

import os
import re
from pathlib import Path

# 尝试导入所需的库
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告：未安装 python-docx 库，Word 转换功能将不可用。请运行: pip install python-docx")

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False
    print("警告：未安装 reportlab 库，PDF 转换功能将不可用。请运行: pip install reportlab")


def parse_text_file(txt_path):
    """
    解析识别文本文件
    
    Args:
        txt_path: 文本文件路径
    
    Returns:
        dict: 包含元数据和对话列表的字典
    """
    with open(txt_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 解析元数据
    metadata = {}
    content_start = 0
    
    for i, line in enumerate(lines):
        line = line.strip()
        if line.startswith('识别时间:'):
            metadata['识别时间'] = line.replace('识别时间:', '').strip()
        elif line.startswith('音频URL:'):
            metadata['音频URL'] = line.replace('音频URL:', '').strip()
        elif line.startswith('说话人识别:'):
            metadata['说话人识别'] = line.replace('说话人识别:', '').strip()
        elif '=' in line and len(line) > 20:  # 分隔线
            content_start = i + 1
            break
    
    # 解析对话内容
    dialogues = []
    pattern = r'\[([^\]]+)\]\s*(说话人[A-Z])\s*(.+)'
    
    for line in lines[content_start:]:
        line = line.strip()
        if not line:
            continue
        
        match = re.match(pattern, line)
        if match:
            time_range = match.group(1)
            speaker = match.group(2)
            content = match.group(3)
            dialogues.append({
                'time': time_range,
                'speaker': speaker,
                'content': content
            })
    
    return {
        'metadata': metadata,
        'dialogues': dialogues
    }


def convert_to_word(txt_path, output_path=None):
    """
    转换为 Word 格式
    
    Args:
        txt_path: 输入文本文件路径
        output_path: 输出 Word 文件路径（可选）
    
    Returns:
        str: 输出文件路径
    """
    if not HAS_DOCX:
        raise ImportError("未安装 python-docx 库，请运行: pip install python-docx")
    
    # 解析文本文件
    data = parse_text_file(txt_path)
    
    # 确定输出路径
    if output_path is None:
        base_name = os.path.splitext(txt_path)[0]
        output_path = f"{base_name}.docx"
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 添加标题
    title = doc.add_heading('语音识别文本', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加元数据
    if data['metadata']:
        doc.add_paragraph()  # 空行
        for key, value in data['metadata'].items():
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(value)
    
    # 添加分隔线
    doc.add_paragraph('=' * 50)
    doc.add_paragraph()  # 空行
    
    # 添加对话内容
    for dialogue in data['dialogues']:
        # 时间和说话人（加粗、蓝色）
        time_speaker = doc.add_paragraph()
        run = time_speaker.add_run(f"[{dialogue['time']}] {dialogue['speaker']}")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
        
        # 说话内容
        content_para = doc.add_paragraph(dialogue['content'])
        content_para.paragraph_format.left_indent = Pt(20)  # 缩进
        
        doc.add_paragraph()  # 空行
    
    # 保存文档
    doc.save(output_path)
    return output_path


def convert_to_pdf(txt_path, output_path=None):
    """
    转换为 PDF 格式
    
    Args:
        txt_path: 输入文本文件路径
        output_path: 输出 PDF 文件路径（可选）
    
    Returns:
        str: 输出文件路径
    """
    if not HAS_REPORTLAB:
        raise ImportError("未安装 reportlab 库，请运行: pip install reportlab")
    
    # 解析文本文件
    data = parse_text_file(txt_path)
    
    # 确定输出路径
    if output_path is None:
        base_name = os.path.splitext(txt_path)[0]
        output_path = f"{base_name}.pdf"
    
    # 创建 PDF 文档
    doc = SimpleDocTemplate(output_path, pagesize=A4)
    story = []
    
    # 获取样式
    styles = getSampleStyleSheet()
    
    # 创建自定义样式
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=12,
        alignment=1  # 居中
    )
    
    metadata_style = ParagraphStyle(
        'CustomMetadata',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6
    )
    
    time_speaker_style = ParagraphStyle(
        'CustomTimeSpeaker',
        parent=styles['Normal'],
        fontSize=11,
        textColor=HexColor('#0000FF'),  # 蓝色
        spaceAfter=3,
        fontName='Helvetica-Bold'
    )
    
    content_style = ParagraphStyle(
        'CustomContent',
        parent=styles['Normal'],
        fontSize=11,
        leftIndent=20,
        spaceAfter=6
    )
    
    # 添加标题
    story.append(Paragraph('语音识别文本', title_style))
    story.append(Spacer(1, 0.2 * inch))
    
    # 添加元数据
    if data['metadata']:
        for key, value in data['metadata'].items():
            text = f"<b>{key}:</b> {value}"
            story.append(Paragraph(text, metadata_style))
        story.append(Spacer(1, 0.1 * inch))
    
    # 添加分隔线
    story.append(Paragraph('=' * 50, metadata_style))
    story.append(Spacer(1, 0.1 * inch))
    
    # 添加对话内容
    for dialogue in data['dialogues']:
        # 时间和说话人
        time_speaker_text = f"[{dialogue['time']}] {dialogue['speaker']}"
        story.append(Paragraph(time_speaker_text, time_speaker_style))
        
        # 说话内容（转义 HTML 特殊字符）
        content = dialogue['content'].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        story.append(Paragraph(content, content_style))
        
        story.append(Spacer(1, 0.05 * inch))
    
    # 构建 PDF
    doc.build(story)
    return output_path


def main():
    """命令行入口"""
    import sys
    
    if len(sys.argv) < 2:
        print("使用方法: python convert_to_doc.py <文本文件路径> [输出格式: pdf|docx] [输出路径]")
        print("\n示例:")
        print("  python convert_to_doc.py 识别文本_20260108_032840.txt")
        print("  python convert_to_doc.py 识别文本_20260108_032840.txt pdf")
        print("  python convert_to_doc.py 识别文本_20260108_032840.txt docx output.docx")
        sys.exit(1)
    
    txt_path = sys.argv[1]
    if not os.path.exists(txt_path):
        print(f"错误：文件不存在: {txt_path}")
        sys.exit(1)
    
    output_format = sys.argv[2] if len(sys.argv) > 2 else 'docx'
    output_path = sys.argv[3] if len(sys.argv) > 3 else None
    
    try:
        if output_format.lower() == 'pdf':
            result_path = convert_to_pdf(txt_path, output_path)
            print(f"✅ PDF 文件已生成: {result_path}")
        elif output_format.lower() == 'docx':
            result_path = convert_to_word(txt_path, output_path)
            print(f"✅ Word 文件已生成: {result_path}")
        else:
            print(f"错误：不支持的格式 '{output_format}'，请使用 'pdf' 或 'docx'")
            sys.exit(1)
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()


"""
将识别文本转换为 PDF 或 Word 格式
支持格式调整：说话人和时间放在上面一行，说话内容放在下面一行
"""

import os
import re
from pathlib import Path

# 尝试导入所需的库
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告：未安装 python-docx 库，Word 转换功能将不可用。请运行: pip install python-docx")

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False
    print("警告：未安装 reportlab 库，PDF 转换功能将不可用。请运行: pip install reportlab")


def parse_text_file(txt_path):
    """
    解析识别文本文件
    
    Args:
        txt_path: 文本文件路径
    
    Returns:
        dict: 包含元数据和对话列表的字典
    """
    with open(txt_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 解析元数据
    metadata = {}
    content_start = 0
    
    for i, line in enumerate(lines):
        line = line.strip()
        if line.startswith('识别时间:'):
            metadata['识别时间'] = line.replace('识别时间:', '').strip()
        elif line.startswith('音频URL:'):
            metadata['音频URL'] = line.replace('音频URL:', '').strip()
        elif line.startswith('说话人识别:'):
            metadata['说话人识别'] = line.replace('说话人识别:', '').strip()
        elif '=' in line and len(line) > 20:  # 分隔线
            content_start = i + 1
            break
    
    # 解析对话内容
    dialogues = []
    pattern = r'\[([^\]]+)\]\s*(说话人[A-Z])\s*(.+)'
    
    for line in lines[content_start:]:
        line = line.strip()
        if not line:
            continue
        
        match = re.match(pattern, line)
        if match:
            time_range = match.group(1)
            speaker = match.group(2)
            content = match.group(3)
            dialogues.append({
                'time': time_range,
                'speaker': speaker,
                'content': content
            })
    
    return {
        'metadata': metadata,
        'dialogues': dialogues
    }


def convert_to_word(txt_path, output_path=None):
    """
    转换为 Word 格式
    
    Args:
        txt_path: 输入文本文件路径
        output_path: 输出 Word 文件路径（可选）
    
    Returns:
        str: 输出文件路径
    """
    if not HAS_DOCX:
        raise ImportError("未安装 python-docx 库，请运行: pip install python-docx")
    
    # 解析文本文件
    data = parse_text_file(txt_path)
    
    # 确定输出路径
    if output_path is None:
        base_name = os.path.splitext(txt_path)[0]
        output_path = f"{base_name}.docx"
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 添加标题
    title = doc.add_heading('语音识别文本', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加元数据
    if data['metadata']:
        doc.add_paragraph()  # 空行
        for key, value in data['metadata'].items():
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(value)
    
    # 添加分隔线
    doc.add_paragraph('=' * 50)
    doc.add_paragraph()  # 空行
    
    # 添加对话内容
    for dialogue in data['dialogues']:
        # 时间和说话人（加粗、蓝色）
        time_speaker = doc.add_paragraph()
        run = time_speaker.add_run(f"[{dialogue['time']}] {dialogue['speaker']}")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
        
        # 说话内容
        content_para = doc.add_paragraph(dialogue['content'])
        content_para.paragraph_format.left_indent = Pt(20)  # 缩进
        
        doc.add_paragraph()  # 空行
    
    # 保存文档
    doc.save(output_path)
    return output_path


def convert_to_pdf(txt_path, output_path=None):
    """
    转换为 PDF 格式
    
    Args:
        txt_path: 输入文本文件路径
        output_path: 输出 PDF 文件路径（可选）
    
    Returns:
        str: 输出文件路径
    """
    if not HAS_REPORTLAB:
        raise ImportError("未安装 reportlab 库，请运行: pip install reportlab")
    
    # 解析文本文件
    data = parse_text_file(txt_path)
    
    # 确定输出路径
    if output_path is None:
        base_name = os.path.splitext(txt_path)[0]
        output_path = f"{base_name}.pdf"
    
    # 创建 PDF 文档
    doc = SimpleDocTemplate(output_path, pagesize=A4)
    story = []
    
    # 获取样式
    styles = getSampleStyleSheet()
    
    # 创建自定义样式
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=12,
        alignment=1  # 居中
    )
    
    metadata_style = ParagraphStyle(
        'CustomMetadata',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6
    )
    
    time_speaker_style = ParagraphStyle(
        'CustomTimeSpeaker',
        parent=styles['Normal'],
        fontSize=11,
        textColor=HexColor('#0000FF'),  # 蓝色
        spaceAfter=3,
        fontName='Helvetica-Bold'
    )
    
    content_style = ParagraphStyle(
        'CustomContent',
        parent=styles['Normal'],
        fontSize=11,
        leftIndent=20,
        spaceAfter=6
    )
    
    # 添加标题
    story.append(Paragraph('语音识别文本', title_style))
    story.append(Spacer(1, 0.2 * inch))
    
    # 添加元数据
    if data['metadata']:
        for key, value in data['metadata'].items():
            text = f"<b>{key}:</b> {value}"
            story.append(Paragraph(text, metadata_style))
        story.append(Spacer(1, 0.1 * inch))
    
    # 添加分隔线
    story.append(Paragraph('=' * 50, metadata_style))
    story.append(Spacer(1, 0.1 * inch))
    
    # 添加对话内容
    for dialogue in data['dialogues']:
        # 时间和说话人
        time_speaker_text = f"[{dialogue['time']}] {dialogue['speaker']}"
        story.append(Paragraph(time_speaker_text, time_speaker_style))
        
        # 说话内容（转义 HTML 特殊字符）
        content = dialogue['content'].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        story.append(Paragraph(content, content_style))
        
        story.append(Spacer(1, 0.05 * inch))
    
    # 构建 PDF
    doc.build(story)
    return output_path


def main():
    """命令行入口"""
    import sys
    
    if len(sys.argv) < 2:
        print("使用方法: python convert_to_doc.py <文本文件路径> [输出格式: pdf|docx] [输出路径]")
        print("\n示例:")
        print("  python convert_to_doc.py 识别文本_20260108_032840.txt")
        print("  python convert_to_doc.py 识别文本_20260108_032840.txt pdf")
        print("  python convert_to_doc.py 识别文本_20260108_032840.txt docx output.docx")
        sys.exit(1)
    
    txt_path = sys.argv[1]
    if not os.path.exists(txt_path):
        print(f"错误：文件不存在: {txt_path}")
        sys.exit(1)
    
    output_format = sys.argv[2] if len(sys.argv) > 2 else 'docx'
    output_path = sys.argv[3] if len(sys.argv) > 3 else None
    
    try:
        if output_format.lower() == 'pdf':
            result_path = convert_to_pdf(txt_path, output_path)
            print(f"✅ PDF 文件已生成: {result_path}")
        elif output_format.lower() == 'docx':
            result_path = convert_to_word(txt_path, output_path)
            print(f"✅ Word 文件已生成: {result_path}")
        else:
            print(f"错误：不支持的格式 '{output_format}'，请使用 'pdf' 或 'docx'")
            sys.exit(1)
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()


"""
将识别文本转换为 PDF 或 Word 格式
支持格式调整：说话人和时间放在上面一行，说话内容放在下面一行
"""

import os
import re
from pathlib import Path

# 尝试导入所需的库
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告：未安装 python-docx 库，Word 转换功能将不可用。请运行: pip install python-docx")

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False
    print("警告：未安装 reportlab 库，PDF 转换功能将不可用。请运行: pip install reportlab")


def parse_text_file(txt_path):
    """
    解析识别文本文件
    
    Args:
        txt_path: 文本文件路径
    
    Returns:
        dict: 包含元数据和对话列表的字典
    """
    with open(txt_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 解析元数据
    metadata = {}
    content_start = 0
    
    for i, line in enumerate(lines):
        line = line.strip()
        if line.startswith('识别时间:'):
            metadata['识别时间'] = line.replace('识别时间:', '').strip()
        elif line.startswith('音频URL:'):
            metadata['音频URL'] = line.replace('音频URL:', '').strip()
        elif line.startswith('说话人识别:'):
            metadata['说话人识别'] = line.replace('说话人识别:', '').strip()
        elif '=' in line and len(line) > 20:  # 分隔线
            content_start = i + 1
            break
    
    # 解析对话内容
    dialogues = []
    pattern = r'\[([^\]]+)\]\s*(说话人[A-Z])\s*(.+)'
    
    for line in lines[content_start:]:
        line = line.strip()
        if not line:
            continue
        
        match = re.match(pattern, line)
        if match:
            time_range = match.group(1)
            speaker = match.group(2)
            content = match.group(3)
            dialogues.append({
                'time': time_range,
                'speaker': speaker,
                'content': content
            })
    
    return {
        'metadata': metadata,
        'dialogues': dialogues
    }


def convert_to_word(txt_path, output_path=None):
    """
    转换为 Word 格式
    
    Args:
        txt_path: 输入文本文件路径
        output_path: 输出 Word 文件路径（可选）
    
    Returns:
        str: 输出文件路径
    """
    if not HAS_DOCX:
        raise ImportError("未安装 python-docx 库，请运行: pip install python-docx")
    
    # 解析文本文件
    data = parse_text_file(txt_path)
    
    # 确定输出路径
    if output_path is None:
        base_name = os.path.splitext(txt_path)[0]
        output_path = f"{base_name}.docx"
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 添加标题
    title = doc.add_heading('语音识别文本', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加元数据
    if data['metadata']:
        doc.add_paragraph()  # 空行
        for key, value in data['metadata'].items():
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(value)
    
    # 添加分隔线
    doc.add_paragraph('=' * 50)
    doc.add_paragraph()  # 空行
    
    # 添加对话内容
    for dialogue in data['dialogues']:
        # 时间和说话人（加粗、蓝色）
        time_speaker = doc.add_paragraph()
        run = time_speaker.add_run(f"[{dialogue['time']}] {dialogue['speaker']}")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
        
        # 说话内容
        content_para = doc.add_paragraph(dialogue['content'])
        content_para.paragraph_format.left_indent = Pt(20)  # 缩进
        
        doc.add_paragraph()  # 空行
    
    # 保存文档
    doc.save(output_path)
    return output_path


def convert_to_pdf(txt_path, output_path=None):
    """
    转换为 PDF 格式
    
    Args:
        txt_path: 输入文本文件路径
        output_path: 输出 PDF 文件路径（可选）
    
    Returns:
        str: 输出文件路径
    """
    if not HAS_REPORTLAB:
        raise ImportError("未安装 reportlab 库，请运行: pip install reportlab")
    
    # 解析文本文件
    data = parse_text_file(txt_path)
    
    # 确定输出路径
    if output_path is None:
        base_name = os.path.splitext(txt_path)[0]
        output_path = f"{base_name}.pdf"
    
    # 创建 PDF 文档
    doc = SimpleDocTemplate(output_path, pagesize=A4)
    story = []
    
    # 获取样式
    styles = getSampleStyleSheet()
    
    # 创建自定义样式
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=12,
        alignment=1  # 居中
    )
    
    metadata_style = ParagraphStyle(
        'CustomMetadata',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6
    )
    
    time_speaker_style = ParagraphStyle(
        'CustomTimeSpeaker',
        parent=styles['Normal'],
        fontSize=11,
        textColor=HexColor('#0000FF'),  # 蓝色
        spaceAfter=3,
        fontName='Helvetica-Bold'
    )
    
    content_style = ParagraphStyle(
        'CustomContent',
        parent=styles['Normal'],
        fontSize=11,
        leftIndent=20,
        spaceAfter=6
    )
    
    # 添加标题
    story.append(Paragraph('语音识别文本', title_style))
    story.append(Spacer(1, 0.2 * inch))
    
    # 添加元数据
    if data['metadata']:
        for key, value in data['metadata'].items():
            text = f"<b>{key}:</b> {value}"
            story.append(Paragraph(text, metadata_style))
        story.append(Spacer(1, 0.1 * inch))
    
    # 添加分隔线
    story.append(Paragraph('=' * 50, metadata_style))
    story.append(Spacer(1, 0.1 * inch))
    
    # 添加对话内容
    for dialogue in data['dialogues']:
        # 时间和说话人
        time_speaker_text = f"[{dialogue['time']}] {dialogue['speaker']}"
        story.append(Paragraph(time_speaker_text, time_speaker_style))
        
        # 说话内容（转义 HTML 特殊字符）
        content = dialogue['content'].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        story.append(Paragraph(content, content_style))
        
        story.append(Spacer(1, 0.05 * inch))
    
    # 构建 PDF
    doc.build(story)
    return output_path


def main():
    """命令行入口"""
    import sys
    
    if len(sys.argv) < 2:
        print("使用方法: python convert_to_doc.py <文本文件路径> [输出格式: pdf|docx] [输出路径]")
        print("\n示例:")
        print("  python convert_to_doc.py 识别文本_20260108_032840.txt")
        print("  python convert_to_doc.py 识别文本_20260108_032840.txt pdf")
        print("  python convert_to_doc.py 识别文本_20260108_032840.txt docx output.docx")
        sys.exit(1)
    
    txt_path = sys.argv[1]
    if not os.path.exists(txt_path):
        print(f"错误：文件不存在: {txt_path}")
        sys.exit(1)
    
    output_format = sys.argv[2] if len(sys.argv) > 2 else 'docx'
    output_path = sys.argv[3] if len(sys.argv) > 3 else None
    
    try:
        if output_format.lower() == 'pdf':
            result_path = convert_to_pdf(txt_path, output_path)
            print(f"✅ PDF 文件已生成: {result_path}")
        elif output_format.lower() == 'docx':
            result_path = convert_to_word(txt_path, output_path)
            print(f"✅ Word 文件已生成: {result_path}")
        else:
            print(f"错误：不支持的格式 '{output_format}'，请使用 'pdf' 或 'docx'")
            sys.exit(1)
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()

